from rest_framework import generics, permissions, status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from django.contrib.auth import get_user_model
from openpyxl import load_workbook
from decimal import Decimal
import csv
from django.http import HttpResponse, JsonResponse
import xlsxwriter
from io import BytesIO
from datetime import datetime
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.negotiation import BaseContentNegotiation

from rest_framework.response import Response

from django.utils.dateparse import parse_date
from openpyxl import Workbook
from docx import Document
from docx.shared import Inches
from datetime import datetime
import io
from .models import (
    School, TransferReceived, Distribution, Report
)
from .serializers import (
    SchoolSerializer, TransferReceivedSerializer,
    DistributionSerializer, ReportSerializer, AdminUserSerializer,
    MyTokenObtainPairSerializer, 
)
from .utils import generate_pdf_report
from rest_framework import status
from django.contrib.auth import authenticate
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework_simplejwt.views import TokenObtainPairView
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework import permissions, generics
from django.shortcuts import get_object_or_404
from rest_framework.renderers import BaseRenderer
from django.http import FileResponse
from rest_framework.parsers import JSONParser

document = Document()
document.add_heading("Dusangire Lunch Report", level=1)
document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
document.save("report.docx")

@api_view(["PUT"])
@permission_classes([IsAuthenticated])
def mark_transfer_as_pending_delete(request, pk):
    try:
        transfer = TransferReceived.objects.get(pk=pk)
    except TransferReceived.DoesNotExist:
        return Response({"error": "Transfer not found"}, status=404)

    user = request.user
    role = check_user_role(request)
    if role["role"] not in ["admin", "superuser"]:
        return Response({"error": "Only admins or superusers can request deletion."}, status=403)


    reason = request.data.get("delete_reason", "")
    transfer.delete_status = "pending"
    transfer.delete_reason = reason
    transfer.save()
    return Response({"message": "Marked as pending deletion."})


class MyTokenObtainPairView(TokenObtainPairView):
    serializer_class = MyTokenObtainPairSerializer



User = get_user_model()

# ✅ Role check helper
def check_user_role(request):
    user = request.user
    if not user.is_authenticated:
        return {"status": "error", "message": "Authentication required", "code": 401}
    if user.is_superuser:
        return {"status": "ok", "role": "superuser"}
    elif user.is_staff:
        return {"status": "ok", "role": "admin"}
    else:
        return {"status": "error", "message": "Unauthorized - Not admin", "code": 403}

# ✅ API view to return only the role
@api_view(["GET"])
@permission_classes([IsAuthenticated])
def check_user_role_view(request):
    user = request.user
    if user.is_superuser:
        return Response({"role": "superadmin"})
    elif user.is_staff:
        return Response({"role": "admin"})
    return Response({"role": "unauthorized"}, status=403)

# ✅ Health Check
def health_check(request):
    return JsonResponse({"status": "healthy"})


def is_superuser_or_403(request):
    if not request.user.is_authenticated:
        return Response({"detail": "Authentication required"}, status=status.HTTP_401_UNAUTHORIZED)
    if not request.user.is_superuser:
        return Response({"detail": "You do not have permission to perform this action"}, status=status.HTTP_403_FORBIDDEN)
    return None  # means allowed
# ✅ School Views

class SchoolListCreateAPIView(generics.ListCreateAPIView):
    queryset = School.objects.filter(delete_status__in=['active', 'pending']).order_by('-created_at')
    serializer_class = SchoolSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]

# 🔁 Mark school as pending delete
class MarkSchoolPendingDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def put(self, request, pk):
        school = get_object_or_404(School, pk=pk)
        reason = request.data.get("delete_reason")

        if not reason:
            return Response({"detail": "Reason for deletion is required."}, status=status.HTTP_400_BAD_REQUEST)

        school.delete_status = "pending"
        school.delete_reason = reason
        school.save()
        return Response({"detail": "School marked as pending deletion."})

# ✅ Recover a pending-deleted school
class RecoverSchoolView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def put(self, request, pk):
        school = get_object_or_404(School, pk=pk)

        if school.delete_status != "pending":
            return Response({"detail": "Only schools with pending deletion can be recovered."}, status=400)

        school.delete_status = "active"
        school.delete_reason = ""
        school.save()
        return Response({"detail": "School successfully recovered."})

# ❌ Confirm deletion by superuser
class ConfirmSchoolDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, pk):
        school = get_object_or_404(School, pk=pk)

        if not request.user.is_superuser:
            return Response({"detail": "Only superusers can confirm deletion."}, status=403)

        if school.delete_status != "pending":
            return Response({"detail": "School is not marked for deletion."}, status=400)

        school.delete_status = "deleted"
        school.save()
        return Response({"detail": "School permanently deleted (soft)."})
# ✅ Transfer Views
class TransferReceivedListCreateAPIView(generics.ListCreateAPIView):
    serializer_class = TransferReceivedSerializer
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        return TransferReceived.objects.filter(delete_status__in=['active', 'pending']).order_by('-timestamp')

class ConfirmTransferDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, pk):
        try:
            transfer = TransferReceived.objects.get(pk=pk)
        except TransferReceived.DoesNotExist:
            return Response({"error": "Transfer not found."}, status=404)
                              
        role = check_user_role(request)
        if role["status"] != "ok" or role["role"] != "superuser":
            return Response({"error": "Only superusers can confirm deletions."}, status=403)

        if transfer.delete_status != "pending":
            return Response({"error": "Transfer is not marked for deletion."}, status=400)

        transfer.delete_status = "deleted"
        transfer.save()
        return Response({"message": "Transfer marked as deleted."}, status=200)



class TransferExcelUploadView(APIView):
    parser_classes = [MultiPartParser, FormParser]
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        role = check_user_role(request)
        if role["status"] != "ok":
            return Response({"detail": role["message"]}, status=role["code"])
        if role["role"] != "superuser":
            return Response({"detail": "Only superusers can upload transfers via Excel."}, status=403)

        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'error': 'No file uploaded.'}, status=400)

        try:
            workbook = load_workbook(file_obj)
            sheet = workbook.active

            headers = [cell.value for cell in sheet[1]]
            col_index = {header.strip().lower(): idx for idx, header in enumerate(headers)}

            required_columns = ['school_code', 'donor', 'amount', 'school_name']
            for col in required_columns:
                if col not in col_index:
                    return Response({'error': f'Missing required column: {col}'}, status=400)

            for row in sheet.iter_rows(min_row=2, values_only=True):
                school_code = row[col_index['school_code']]
                donor = row[col_index['donor']]
                amount = row[col_index['amount']]
                school_name = row[col_index['school_name']]
                number_of_transactions = row[col_index.get('number_of_transactions', -1)]
                contribution_type = row[col_index.get('contribution_type', -1)]

                # Make sure we handle missing optional fields
                account_number = account_number if account_number else ""
                number_of_transactions = int(number_of_transactions) if number_of_transactions else 0
                contribution_type = contribution_type if contribution_type else None

                # Find all matching schools
                matching_schools = School.objects.filter(name__iexact=school_name)
                # Set AccountNumber on school(s) if present
                if matching_schools.exists() and account_number:
                    for school in matching_schools:
                        school.AccountNumber = account_number
                        school.save()

                transfer = TransferReceived.objects.create(
                    SchoolCode=school_code,
                    Donor=donor,
                    Amount=Decimal(amount),
                    NumberOfTransactions=number_of_transactions,
                    contribution_type=contribution_type
                )

                if matching_schools.exists():
                    transfer.SchoolName.set(matching_schools)

            return Response({'message': 'Transfers uploaded successfully.'}, status=201)

        except Exception as e:
            return Response({'error': str(e)}, status=500)
# ✅ Distribution Views
class DistributionListCreateAPIView(generics.ListCreateAPIView):
    queryset = Distribution.objects.filter(delete_status__in=["active", "pending"]).order_by('-distributed_on')
    serializer_class = DistributionSerializer
    permission_classes = [IsAuthenticated]

class DistributionListAPIView(generics.ListAPIView):
    queryset = Distribution.objects.filter(delete_status__in=["active", "pending"]).order_by('-distributed_on')
    serializer_class = DistributionSerializer

class MarkDistributionPendingDelete(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def put(self, request, pk):
        try:
            dist = Distribution.objects.get(pk=pk)
        except Distribution.DoesNotExist:
            return Response({"error": "Distribution not found."}, status=404)

        role = check_user_role(request)
        if role["status"] != "ok":
            return Response({"detail": role["message"]}, status=role["code"])

        if role["role"] not in ["admin", "superuser"]:
            return Response({"error": "Only admins can request deletion."}, status=403)

        reason = request.data.get("delete_reason")
        if not reason:
            return Response({"error": "Delete reason is required."}, status=400)

        dist.delete_status = "pending"
        dist.delete_reason = reason
        dist.save()

        return Response({"message": "Marked as pending deletion."}, status=200)


class ConfirmDistributionDeletion(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, pk):
        try:
            dist = Distribution.objects.get(pk=pk)
        except Distribution.DoesNotExist:
            return Response({"error": "Distribution not found."}, status=404)

        role = check_user_role(request)
        if role["status"] != "ok" or role["role"] != "superuser":
            return Response({"error": "Only superusers can confirm deletions."}, status=403)

        if dist.delete_status != "pending":
            return Response({"error": "Distribution is not marked for deletion."}, status=400)

        dist.delete_status = "deleted"
        dist.save()
        return Response({"message": "Distribution marked as deleted."}, status=200)


class RecoverDistribution(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def put(self, request, pk):
        try:
            dist = Distribution.objects.get(pk=pk)
        except Distribution.DoesNotExist:
            return Response({"error": "Distribution not found."}, status=404)

        role = check_user_role(request)
        if role["status"] != "ok" or role["role"] not in ["admin", "superuser"]:
            return Response({"error": "Not allowed to recover."}, status=403)

        dist.delete_status = "active"
        dist.delete_reason = ""
        dist.save()
        return Response({"message": "Distribution recovered."}, status=200)

# ✅ Admin Users
class AdminUserListAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        role = check_user_role(request)
        if role["status"] != "ok":
            return Response({"detail": role["message"]}, status=role["code"])
        if role["role"] != "superuser":
            return Response({"detail": "Only superusers can view all admin users."}, status=403)
        admins = User.objects.all()
        serializer = AdminUserSerializer(admins, many=True)
        return Response(serializer.data)

# ✅ Simulate MoMo Payment
class SimulatePaymentAPIView(APIView):
    def post(self, request):
        serializer = TransferReceivedSerializer(data=request.data)
        if serializer.is_valid():
            transfer = serializer.save()
            return Response({
                'message': 'Payment simulated successfully.',
                'data': TransferReceivedSerializer(transfer).data
            }, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# ✅ Report Views
class ReportListCreateAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        role = check_user_role(request)
        if role["status"] != "ok":
            return Response({"detail": role["message"]}, status=role["code"])
        if role["role"] != "superuser":
            return Response({"detail": "Only superusers can view reports."}, status=403)
        reports = Report.objects.all().order_by('-generated_on')
        serializer = ReportSerializer(reports, many=True)
        return Response(serializer.data)

    def post(self, request):
        role = check_user_role(request)
        if role["status"] != "ok":
            return Response({"detail": role["message"]}, status=role["code"])
        if role["role"] != "superuser":
            return Response({"detail": "Only superusers can create reports."}, status=403)
        serializer = ReportSerializer(data=request.data)
        if serializer.is_valid():
            report = serializer.save()
            return Response(ReportSerializer(report).data, status=201)
        return Response(serializer.errors, status=400)

# ✅ Generate Report View
def check_user_role(request):
    user = request.user
    if not user.is_authenticated:
        return {"status": "fail", "message": "Authentication required", "code": 401}
    if user.is_superuser:
        return {"status": "ok", "role": "superuser"}
    elif user.is_staff:
        return {"status": "ok", "role": "admin"}
    else:
        return {"status": "ok", "role": "user"}

class IgnoreClientContentNegotiation(BaseContentNegotiation):
    def select_renderer(self, request, renderers, format_suffix):
        return (None, None)
    def select_parser(self, request, parsers):
        for parser in parsers:
            if isinstance(parser, JSONParser):
                return parser
        return parsers[0]

class GenerateReportView(APIView):
    permission_classes = [IsAuthenticated]
    content_negotiation_class = IgnoreClientContentNegotiation
    parser_classes = [JSONParser]

    def post(self, request):
        role = check_user_role(request)
        if role["status"] != "ok":
            return Response({"detail": role["message"]}, status=role["code"])
        if role["role"] != "superuser":
            return Response({"detail": "Only superusers can generate reports."}, status=403)

        report_type = request.data.get("report_type")
        format = request.data.get("format", "excel")
        selected_columns = request.data.get("columns", [])
        selected_rows = request.data.get("selected_rows", [])

        # If selected_rows are provided, skip queryset and use them directly
        if selected_rows:
            try:
                if format == "excel":
                    return self.generate_excel_from_rows(selected_rows, selected_columns, report_type)
                elif format == "word":
                    return self.generate_word_from_rows(selected_rows, selected_columns, report_type)
                else:
                    return Response({"detail": "Unsupported format"}, status=400)
            except Exception as e:
                return Response({"detail": f"Error generating report: {str(e)}"}, status=500)

        # Fallback to queryset (no selected rows sent)
        try:
            queryset = self.get_queryset(report_type, request)
        except Exception as e:
            return Response({"detail": f"Error filtering data: {str(e)}"}, status=400)

        try:
            if format == "excel":
                return self.generate_excel(queryset, report_type, selected_columns)
            elif format == "word":
                return self.generate_word(queryset, report_type, selected_columns)
            else:
                return Response({"detail": "Unsupported format"}, status=400)
        except Exception as e:
            return Response({"detail": f"Error generating report: {str(e)}"}, status=500)

    def get_queryset(self, report_type, request):
        start_date = parse_date(request.data.get("start_date")) if request.data.get("start_date") else None
        end_date = parse_date(request.data.get("end_date")) if request.data.get("end_date") else None
        school_code = request.data.get("school_code")

        if report_type == "contributions":
            queryset = TransferReceived.objects.all()
            if school_code:
                queryset = queryset.filter(SchoolCode=school_code)
            if start_date and end_date:
                queryset = queryset.filter(timestamp__date__range=(start_date, end_date))

        elif report_type == "distributions":
            queryset = Distribution.objects.all()
            if school_code:
                queryset = queryset.filter(school__name=school_code)
            if start_date and end_date:
                queryset = queryset.filter(distributed_on__date__range=(start_date, end_date))

        elif report_type == "schools":
            queryset = School.objects.all()
            if school_code:
                queryset = queryset.filter(name=school_code)
        else:
            raise ValueError("Invalid report type")

        return queryset

    def generate_excel_from_rows(self, rows, selected_columns, report_type):
        wb = Workbook()
        ws = wb.active
        ws.title = "Report"

        headers = selected_columns or list(rows[0].keys())
        ws.append(headers)

        for row in rows:
            ws.append([row.get(col, "") for col in headers])

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return FileResponse(output, as_attachment=True, filename=f"{report_type}_selected_rows.xlsx")

    def generate_word_from_rows(self, rows, selected_columns, report_type):
        doc = Document()
        doc.add_heading(f"{report_type.capitalize()} Report", 0)

        headers = selected_columns or list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        for i, col in enumerate(headers):
            table.rows[0].cells[i].text = col

        for row_data in rows:
            row = table.add_row().cells
            for i, col in enumerate(headers):
                row[i].text = str(row_data.get(col, "—"))

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return FileResponse(
            output,
            as_attachment=True,
            filename=f"{report_type}_selected_rows.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def generate_excel(self, queryset, report_type, selected_columns):
        wb = Workbook()
        ws = wb.active
        ws.title = "Report"

        field_map = self.get_field_map(report_type)
        fields = field_map or {}
        headers = selected_columns or list(fields.keys())
        ws.append(headers)

        for obj in queryset:
            row = [str(fields[col](obj)) if col in fields else "—" for col in headers]
            ws.append(row)

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return FileResponse(output, as_attachment=True, filename=f"{report_type}_report.xlsx")

    def generate_word(self, queryset, report_type, selected_columns):
        doc = Document()
        doc.add_heading(f"{report_type.capitalize()} Report", 0)

        field_map = self.get_field_map(report_type)
        fields = field_map or {}
        headers = selected_columns or list(fields.keys())

        table = doc.add_table(rows=1, cols=len(headers))
        for i, col in enumerate(headers):
            table.rows[0].cells[i].text = col

        for obj in queryset:
            row = table.add_row().cells
            for i, col in enumerate(headers):
                row[i].text = str(fields[col](obj)) if col in fields else "—"

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return FileResponse(
            output,
            as_attachment=True,
            filename=f"{report_type}_report.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def get_field_map(self, report_type):
        return {
            "contributions": {
                "School Code": lambda obj: obj.SchoolCode,
                "School Name": lambda obj: obj.schools.first().name if obj.schools.exists() else "N/A",
                "Donor": lambda obj: obj.Donor,
                "Amount": lambda obj: obj.Amount,
                "Transactions": lambda obj: obj.NumberOfTransactions,
                "Date": lambda obj: obj.timestamp.strftime("%Y-%m-%d") if obj.timestamp else "N/A"
            },
            "distributions": {
                "School": lambda obj: obj.school.name if obj.school else "N/A",
                "Amount": lambda obj: obj.amount,
                "Date": lambda obj: obj.distributed_on.strftime("%Y-%m-%d") if obj.distributed_on else "N/A"
            },
            "schools": {
                "Name": lambda obj: obj.name,
                "Sector": lambda obj: obj.sector,
                "District": lambda obj: obj.district
            }
        }
# ✅ Transaction Summary View
class TransactionSummaryView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        transfers = TransferReceived.objects.all()
        total_amount = sum(t.amount for t in transfers)
        total_transfers = transfers.count()

        return Response({
            "total_transfers": total_transfers,
            "total_amount": float(total_amount),
        })
class LoginAPIView(APIView):
    def post(self, request):
        username = request.data.get("username")
        password = request.data.get("password")

        if not username or not password:
            return Response({"error": "Username and password are required."}, status=400)

        user = authenticate(username=username, password=password)
        if user is None:
            return Response({"error": "Invalid credentials."}, status=401)

        # Get JWT tokens
        refresh = RefreshToken.for_user(user)
        access_token = str(refresh.access_token)
        refresh_token = str(refresh)

        # Identify role
        if user.is_superuser:
            role = "superuser"
        elif user.is_staff:
            role = "admin"
        else:
            return Response({"error": "Unauthorized role."}, status=403)

        return Response({
            "access": access_token,
            "refresh": refresh_token,
            "role": role,
            "username": user.username
        }, status=200)

class RecoverTransferView(APIView):
    permission_classes = [IsAuthenticated]

    def put(self, request, pk):
        role = check_user_role(request)
        if role["status"] != "ok":
            return Response({"detail": role["message"]}, status=role["code"])

        try:
            transfer = TransferReceived.objects.get(pk=pk)
            if transfer.delete_status != "pending":
                return Response({"detail": "Only pending deletions can be recovered."}, status=400)
            
            transfer.delete_status = "active"
            transfer.save()
            return Response({"message": "Transfer recovered successfully."})
        except TransferReceived.DoesNotExist:
            return Response({"detail": "Transfer not found."}, status=404)
        

@api_view(["GET"])
@permission_classes([IsAuthenticated])
def deleted_transfers_list(request):
    # Check superuser permission
    error_response = is_superuser_or_403(request)
    if error_response:
        return error_response

    deleted_transfers = TransferReceived.objects.filter(delete_status="deleted")
    serializer = TransferReceivedSerializer(deleted_transfers, many=True)
    return Response(serializer.data)


@api_view(["GET"])
@permission_classes([IsAuthenticated])

def deleted_distributions_list(request):
    error_response = is_superuser_or_403(request)
    if error_response:
        return error_response

    deleted_distributions = Distribution.objects.filter(delete_status="deleted")
    serializer = DistributionSerializer(deleted_distributions, many=True)
    return Response(serializer.data)


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def deleted_schools_list(request):
    error_response = is_superuser_or_403(request)
    if error_response:
        return error_response

    deleted_schools = School.objects.filter(delete_status="deleted")
    serializer = SchoolSerializer(deleted_schools, many=True)
    return Response(serializer.data)

from rest_framework import mixins

class TransferReceivedUpdateView(generics.UpdateAPIView):
    queryset = TransferReceived.objects.all()
    serializer_class = TransferReceivedSerializer
    permission_classes = [permissions.IsAuthenticated]

    def put(self, request, pk):
        transfer = self.get_object()
        serializer = self.get_serializer(transfer, data=request.data, partial=False)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data)

    def patch(self, request, pk):
        transfer = self.get_object()
        serializer = self.get_serializer(transfer, data=request.data, partial=True)
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response(serializer.data)